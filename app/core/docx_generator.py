import re
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from app.ocr.models import DocumentElement, ElementType


def create_document_from_elements(
    elements: list[DocumentElement],
    institution_name: str = "TITUS SOLUTIONS EXAM LAB",
    subject: str | None = None,
    class_grade: str | None = None,
    total_marks: int | None = None,
    time_allowed: str | None = None,
    notes: str | None = None,
    answers_markdown: str | None = None,
) -> Document:
    doc = Document()

    # Page Margins & Size: A4, 1-inch margins (WD-18, WD-19)
    section = doc.sections[0]
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    # 1. Header Block (WD-04)
    # Configure centered header
    p_header = doc.add_paragraph()
    p_header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_header.paragraph_format.line_spacing = 1.15

    # Institution Name (WD-16: Arial 16pt Bold)
    run_inst = p_header.add_run(f"{institution_name}\n")
    run_inst.font.name = "Arial"
    run_inst.font.size = Pt(16)
    run_inst.font.bold = True
    run_inst.font.no_proof = True

    # Subject, Grade details (WD-15: Arial 12pt)
    header_lines = []
    if subject:
        header_lines.append(f"Subject: {subject}")
    if class_grade:
        header_lines.append(f"Class/Grade: {class_grade}")
    if header_lines:
        run_details = p_header.add_run(" | ".join(header_lines) + "\n")
        run_details.font.name = "Arial"
        run_details.font.size = Pt(12)
        run_details.font.no_proof = True

    # Marks & Time
    limit_lines = []
    if total_marks is not None:
        limit_lines.append(f"Total Marks: {total_marks}")
    if time_allowed:
        limit_lines.append(f"Time Allowed: {time_allowed}")
    if limit_lines:
        run_limits = p_header.add_run(" | ".join(limit_lines) + "\n")
        run_limits.font.name = "Arial"
        run_limits.font.size = Pt(12)
        run_limits.font.no_proof = True

    # Notes
    if notes:
        run_notes = p_header.add_run(f"Notes: {notes}\n")
        run_notes.font.name = "Arial"
        run_notes.font.size = Pt(11)
        run_notes.font.italic = True
        run_notes.font.no_proof = True

    # Dividers (WD-04 bottom line)
    p_divider = doc.add_paragraph()
    p_divider.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_divider_run = p_divider.add_run("—" * 60)
    p_divider_run.font.name = "Arial"
    p_divider_run.font.size = Pt(10)
    p_divider_run.font.color.rgb = RGBColor(180, 180, 180)
    p_divider_run.font.no_proof = True

    # Helper function to add runs with Red unreadable highlighting (WD-13)
    def add_text_runs(paragraph, text, is_bold=False, is_italic=False, font_size=12):
        # We search for any [UNREADABLE ...] block
        pattern = r"(\[UNREADABLE\b.*?\])"
        parts = re.split(pattern, text)
        for part in parts:
            if not part:
                continue
            run = paragraph.add_run(part)
            run.font.name = "Arial"
            run.font.size = Pt(font_size)
            run.font.bold = is_bold
            run.font.italic = is_italic
            run.font.no_proof = True  # Disable spellcheck (WD-02)

            if part.startswith("[UNREADABLE"):
                run.font.color.rgb = RGBColor(255, 0, 0)  # Red text (WD-13)
                run.font.bold = True

    # Process all Elements in sequence
    for element in elements:
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.15  # WD-17

        # Handle indents based on element type (WD-07, WD-09)
        if element.type == ElementType.sub_question:
            p.paragraph_format.left_indent = Inches(0.5)
        elif element.type == ElementType.mcq_option:
            p.paragraph_format.left_indent = Inches(0.75)
        elif element.type == ElementType.instruction:
            p.paragraph_format.left_indent = Inches(0.25)

        # Style based on element type
        if element.type == ElementType.section_heading:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            # Horizontal rule above
            run_line = p.add_run("—" * 50 + "\n")
            run_line.font.name = "Arial"
            run_line.font.size = Pt(10)
            run_line.font.color.rgb = RGBColor(200, 200, 200)
            run_line.font.no_proof = True

            # Bold 14pt (WD-05, WD-16)
            add_text_runs(p, element.text, is_bold=True, font_size=14)

            # Horizontal rule below
            run_line2 = p.add_run("\n" + "—" * 50)
            run_line2.font.name = "Arial"
            run_line2.font.size = Pt(10)
            run_line2.font.color.rgb = RGBColor(200, 200, 200)
            run_line2.font.no_proof = True

        elif element.type == ElementType.instruction:
            # Italics (WD-12)
            add_text_runs(p, element.text, is_italic=True, font_size=12)

        elif element.type in {ElementType.question, ElementType.sub_question}:
            # Normal question text (WD-06: Question numbers appear exactly as extracted)
            add_text_runs(p, element.text, font_size=12)

            # Mark allocations: right-aligned bold on same line (WD-08)
            if element.mark_allocation:
                p.paragraph_format.tab_stops.add_tab_stop(
                    Inches(6.27), alignment=WD_TAB_ALIGNMENT.RIGHT
                )
                tab_run = p.add_run("\t")
                tab_run.font.no_proof = True
                mark_run = p.add_run(element.mark_allocation)
                mark_run.font.name = "Arial"
                mark_run.font.size = Pt(12)
                mark_run.font.bold = True
                mark_run.font.no_proof = True

        elif element.type == ElementType.mcq_option:
            # Option styling
            add_text_runs(p, element.text, font_size=12)

        elif element.type == ElementType.fill_blank:
            # Fill blank spaces
            add_text_runs(p, element.text, font_size=12)

        elif element.type == ElementType.header:
            # Center and italicize matching headers
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_text_runs(p, element.text, is_italic=True, font_size=11)

        elif element.type == ElementType.match_row:
            # Remove empty paragraph
            p_element = p._p
            p_element.getparent().remove(p_element)

            # Insert borderless two-column table (WD-11, OCR-13)
            table = doc.add_table(rows=1, cols=2)
            table.autofit = False
            table.columns[0].width = Inches(3.0)
            table.columns[1].width = Inches(3.0)
            table.style = 'Normal Table'

            cell_a = table.cell(0, 0)
            cell_b = table.cell(0, 1)
            cell_a.width = Inches(3.0)
            cell_b.width = Inches(3.0)

            p_a = cell_a.paragraphs[0]
            p_a.paragraph_format.line_spacing = 1.15
            add_text_runs(p_a, element.match_column_a or "", font_size=12)

            p_b = cell_b.paragraphs[0]
            p_b.paragraph_format.line_spacing = 1.15
            add_text_runs(p_b, element.match_column_b or "", font_size=12)

        else:
            # Normal paragraph/other text
            add_text_runs(p, element.text, font_size=12)

    # Centred page number footer (WD-20)
    # Adding Page Number fields using word XML elements
    footer = section.footer
    f_p = footer.paragraphs[0]
    f_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    f_p_run = f_p.add_run("Page ")
    f_p_run.font.name = "Arial"
    f_p_run.font.size = Pt(10)
    f_p_run.font.no_proof = True

    # Simple page number XML insertion
    fldSimple = OxmlElement("w:fldSimple")
    fldSimple.set(qn("w:instr"), "PAGE")
    f_p._p.append(fldSimple)

    f_p_run2 = f_p.add_run(" of ")
    f_p_run2.font.name = "Arial"
    f_p_run2.font.size = Pt(10)
    f_p_run2.font.no_proof = True

    fldSimple_num = OxmlElement("w:fldSimple")
    fldSimple_num.set(qn("w:instr"), "NUMPAGES")
    f_p._p.append(fldSimple_num)

    if answers_markdown:
        doc.add_page_break()

        # Add ANSWER KEY heading
        p_ans_head = doc.add_paragraph()
        p_ans_head.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_ans_head = p_ans_head.add_run("ANSWER KEY — AI GENERATED")
        run_ans_head.font.name = "Arial"
        run_ans_head.font.size = Pt(14)
        run_ans_head.font.bold = True
        run_ans_head.font.no_proof = True

        # Add AI generated disclaimer
        p_ans_disc = doc.add_paragraph()
        p_ans_disc.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_ans_disc = p_ans_disc.add_run("Note: Answers below are AI-generated. Please verify before use.")
        run_ans_disc.font.name = "Arial"
        run_ans_disc.font.size = Pt(11)
        run_ans_disc.font.italic = True
        run_ans_disc.font.no_proof = True

        # Write parsed text
        for line in answers_markdown.splitlines():
            stripped = line.strip()
            if not stripped:
                continue
            if "Answers below are AI-generated" in stripped:
                continue

            p_line = doc.add_paragraph()
            p_line.paragraph_format.line_spacing = 1.15
            
            if stripped.startswith("#"):
                h_level = 0
                while h_level < len(stripped) and stripped[h_level] == '#':
                    h_level += 1
                h_text = stripped[h_level:].strip()
                add_text_runs(p_line, h_text, is_bold=True, font_size=13)
            else:
                add_text_runs(p_line, stripped, font_size=12)

    return doc
